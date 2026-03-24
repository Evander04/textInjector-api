from docx import Document
from copy import deepcopy
import json
from pathlib import Path
import os
from io import BytesIO
from datetime import datetime
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE_PATH = os.getenv("TEMPLATE_PATH")
templateArray =['Template Ledger.docx','Template Progress.docx','Template Transcript.docx','Template SAP.docx']
cnaTemplateArray = ['CNA Template Ledger.docx', 'CNA Template Progress.docx', 'CNA Template Transcript.docx', 'CNA Template SAP.docx']
CNA_LEDGER_TOTAL = 3150.0

def _replace_across_runs(runs, key, value):
    """Replace key→value even if the key is split across multiple runs.
    Keeps the original run objects and their formatting."""
    if not runs:
        return

    # Build a full string and a map back to (run_idx, char_idx)
    full = []
    idx_map = []
    for ri, r in enumerate(runs):
        t = r.text or ""
        full.append(t)
        idx_map.extend([(ri, cj) for cj in range(len(t))])
    s = "".join(full)
    if not s or key not in s:
        return

    new_s = s.replace(key, value)

    # Clear all run texts (but keep the runs & formatting)
    for r in runs:
        r.text = ""

    # Re-distribute new_s back into the original runs by their original lengths
    # so formatting sticks to the same spans as much as possible.
    # Count original lengths:
    lens = []
    start = 0
    for ri, r in enumerate(runs):
        # original length of this run was how many idx_map entries had run_idx==ri
        orig_len = sum(1 for pair in idx_map if pair[0] == ri)
        lens.append(orig_len)

    pos = 0
    for ri, r in enumerate(runs):
        take = lens[ri]
        if take == 0:
            continue
        chunk = new_s[pos:pos+take]
        r.text = chunk
        pos += len(chunk)
        if pos >= len(new_s):
            break

    # If replacement made the text longer than total original chars,
    # append overflow to the last run to keep formatting.
    if pos < len(new_s):
        runs[-1].text += new_s[pos:]


def _replace_in_paragraph(paragraph, replacements: dict):
    # First, handle simple same-run replacements to avoid heavy work
    for run in paragraph.runs:
        if not run.text:
            continue
        for k, v in replacements.items():
            if k in run.text:
                run.text = run.text.replace(k, v)

    # Then handle cases where placeholders are split across runs
    for k, v in replacements.items():
        _replace_across_runs(paragraph.runs, k, v)


def _replace_in_cell(cell, replacements: dict):
    for p in cell.paragraphs:
        _replace_in_paragraph(p, replacements)
    # Recursively process nested tables too
    for t in cell.tables:
        for row in t.rows:
            for c in row.cells:
                _replace_in_cell(c, replacements)

def injectTemplate(replacements,type):
    fullname = replacements["@firstName"]+replacements["@middleName"]+replacements["@lastName"]
    print(f"======== creating template {type} for {fullname} =========")
    doc = Document(TEMPLATE_PATH+"/"+templateArray[type])

    # Replace in paragraphs
    for p in doc.paragraphs:
        _replace_in_paragraph(p, replacements)

    # Replace in cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_cell(cell, replacements)


    # Write to memory
    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream


def injectCnaTemplate(replacements, type):
    fullname = replacements["@firstName"] + replacements["@middleName"] + replacements["@lastName"]
    print(f"======== creating CNA template {type} for {fullname} =========")
    doc = Document(TEMPLATE_PATH + "/" + cnaTemplateArray[type])

    for p in doc.paragraphs:
        _replace_in_paragraph(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_cell(cell, replacements)

    if type == 0:
        _fill_cna_ledger_table(doc, replacements)

    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream


def _format_money(value):
    return f"${int(value)}" if float(value).is_integer() else f"${float(value):.2f}"


def _fill_cna_ledger_table(doc, replacements):
    if len(doc.tables) < 9:
        return

    ledger_table = doc.tables[8]
    ledger_rows = replacements.get("__cna_ledger_rows__", [])

    _ensure_cna_ledger_row_capacity(ledger_table, len(ledger_rows))

    for idx in range(1, min(len(ledger_table.rows), len(ledger_rows) + 1)):
        row = ledger_table.rows[idx]
        ledger_row = ledger_rows[idx - 1]
        _write_cna_ledger_cell(row.cells[0], ledger_row.get("date", ""))
        _write_cna_ledger_cell(row.cells[1], ledger_row.get("receipt_number", ""))
        _write_cna_ledger_cell(row.cells[2], ledger_row.get("description", ""))
        _write_cna_ledger_cell(row.cells[3], "", bold=True)
        _write_cna_ledger_cell(row.cells[4], ledger_row.get("credited", ""))
        _write_cna_ledger_cell(row.cells[5], ledger_row.get("old_balance", ""))
        _write_cna_ledger_cell(row.cells[6], ledger_row.get("new_balance", ""), bold=True)

    for idx in range(len(ledger_rows) + 1, len(ledger_table.rows)):
        row = ledger_table.rows[idx]
        for cell_index in range(len(row.cells)):
            _write_cna_ledger_cell(row.cells[cell_index], "", bold=(cell_index in {3, 6}))


def _ensure_cna_ledger_row_capacity(ledger_table, row_count):
    if row_count <= 0 or len(ledger_table.rows) - 1 >= row_count:
        return

    template_tr = ledger_table.rows[-1]._tr
    rows_to_add = row_count - (len(ledger_table.rows) - 1)
    for _ in range(rows_to_add):
        ledger_table._tbl.append(deepcopy(template_tr))


def _write_cna_ledger_cell(cell, value, bold=False):
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in paragraph.runs:
        run.text = ""

    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = value
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

gradeDic = {
    "A+":10,
    "A":9,
    "B":8,
    "C":7,
    "":0
}


def _parse_class_date(value):
    if not value:
        return None

    for fmt in ("%m/%d/%Y", "%m-%d-%Y", "%Y-%m-%d"):
        try:
            return datetime.strptime(str(value).strip(), fmt).date()
        except ValueError:
            continue
    return None


def getFinalGrade(studentModules,studentUnits,classType):
    modules = [gradeDic[m] for m in studentModules]
    units = [gradeDic[u] for u in studentUnits]    
    
    if classType == 1:
        return sum(modules)/ 12
    
    if classType == 2:
        return sum(units) / 8
    
    if classType == 3:
        total = sum(modules) + sum(units)
        return total / 20

    has_modules = any(studentModules)
    has_units = any(studentUnits)

    if has_modules and has_units:
        total = sum(modules) + sum(units)
        return total / (len(studentModules) + len(studentUnits))

    if has_modules:
        return sum(modules) / len(studentModules)

    if has_units:
        return sum(units) / len(studentUnits)

    return 0
        
def parseFinalGrade(grade):
    if grade == 10:
        return "A+"
    
    if grade >= 9:
        return "A"
    
    if grade >=8.5:
        return "B+"
    
    if grade >=8:
        return "B"
    
    if grade >=7.5:
        return "C+"
    
    return "C"


def getGPA(grade):
    if grade >= 9:
        return 4.0
    
    if grade >= 8.5:
        return 3.33
    
    if grade >= 8:
        return 3.0
    
    if grade >= 7.5:
        return 2.33
    
    return 2.0


def getCnaFinalGrade(studentModules):
    modules = [gradeDic[m] for m in studentModules]
    if not modules:
        return 0
    return sum(modules) / len(modules)


def getCnaFinalGradeSAP(studentModules, moduleDates, midpoint):
    midpoint_date = _parse_class_date(midpoint)
    if not midpoint_date:
        return 0

    completed_scores = []
    for grade, module_date in zip(studentModules, moduleDates):
        parsed_date = _parse_class_date(module_date)
        if parsed_date and parsed_date <= midpoint_date and grade in gradeDic:
            completed_scores.append(gradeDic[grade])

    if not completed_scores:
        return 0

    return sum(completed_scores) / len(completed_scores)



def getFinalGradeSAP(studentModules,studentUnits,classType):
    modules = [gradeDic[m] for m in studentModules]
    units = [gradeDic[u] for u in studentUnits]    
    
    if classType == 1:
        return sum(modules[:6])/ 6
    
    if classType == 2:
        return sum(units[:4]) / 4
    
    if classType == 3:
        total = sum(modules[:6]) + sum(units[:4])
        return total / 12

    has_modules = any(studentModules)
    has_units = any(studentUnits)

    if has_modules and has_units:
        total = sum(modules[:6]) + sum(units[:4])
        return total / 12

    if has_modules:
        sample = modules[:6] if len(modules) >= 6 else modules
        return sum(sample) / len(sample) if sample else 0

    if has_units:
        sample = units[:4] if len(units) >= 4 else units
        return sum(sample) / len(sample) if sample else 0

    return 0
    

def insertLedgerValues(replacements,student,classObj):
    
    replacements["@ledate1"] = student.get("receiptDates")[0]
    replacements["@ledate2"] = student.get("receiptDates")[1]
    newBalance = int(classObj.total)-int(classObj.registration)
    replacements["@newb1"] = str(newBalance)

    if classObj.classType == 3 and len(student.get("receiptDates")) > 2: # HHA
        replacements["@pay1"] = "350"
        newBalance = newBalance - 350
        replacements["@newb2"] = str(newBalance)
        replacements["@ledate3"] = student.get("receiptDates")[2] if student.get("receiptDates")[2] else ""
        replacements["@rowV1"] = "3"
        replacements["@rowV2"] = classObj.course + " Tuition"
        replacements["@rowV3"] = "$"+str(newBalance)
        replacements["@rowV4"] = "$0"
        replacements["@rowV5"] = "$"+str(newBalance)
        replacements["@rowV6"] = "$0"
    else: 
        replacements["@pay1"] = str(newBalance)
        replacements["@newb2"] = "0"
        replacements["@ledate3"] = ""
        replacements["@rowV1"] = ""
        replacements["@rowV2"] = ""
        replacements["@rowV3"] = ""
        replacements["@rowV4"] = ""
        replacements["@rowV5"] = ""
        replacements["@rowV6"] = ""


def insertCnaLedgerValues(replacements, student, classObj):
    receipt_dates = student.get("receiptDates") or []
    receipt_amounts = student.get("receiptAmounts") or []
    running_balance = CNA_LEDGER_TOTAL
    ledger_rows = []
    receipt_sequence = 1

    total_rows = max(len(receipt_dates), len(receipt_amounts), 4)

    for idx in range(1, total_rows + 1):
        date_value = receipt_dates[idx - 1] if len(receipt_dates) >= idx else ""
        amount_value = receipt_amounts[idx - 1] if len(receipt_amounts) >= idx else ""

        if idx == 1 and (date_value or amount_value):
            amount_value = "100"
            description = "Registration Fee"
        else:
            if date_value or amount_value:
                description = "Tuition + Supplies" if idx == 2 else "Tuition"
            else:
                description = ""

        if amount_value:
            try:
                payment_value = float(str(amount_value).replace("$", "").replace(",", "").strip())
            except ValueError:
                payment_value = 0.0
        else:
            payment_value = 0.0

        if date_value or amount_value:
            old_balance = running_balance
            running_balance = max(running_balance - payment_value, 0)
            ledger_rows.append({
                "date": date_value,
                "receipt_number": str(receipt_sequence),
                "description": description,
                "credited": _format_money(payment_value) if payment_value else "",
                "old_balance": _format_money(old_balance),
                "new_balance": _format_money(running_balance),
            })
            receipt_sequence += 1
        else:
            ledger_rows.append({
                "date": "",
                "receipt_number": "",
                "description": "",
                "credited": "",
                "old_balance": "",
                "new_balance": "",
            })

    replacements["__cna_ledger_rows__"] = ledger_rows
